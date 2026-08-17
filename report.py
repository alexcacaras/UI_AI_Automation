from docx import Document
from docx.shared import Inches

def build_doc(name, shots):
    doc = Document()
    doc.add_heading(name, level=1)
    for img in shots:
        doc.add_picture(img, width=Inches(6))
    doc.save(f"recordings/docs/{name}/{name}.docx")